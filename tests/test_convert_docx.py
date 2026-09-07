"""python-docx conversion tests (no Docling OCR)."""

from __future__ import annotations

from docx import Document

from scan_to_markdown_docling import _convert_with_python_docx


def test_convert_with_python_docx_heading_and_table(tmp_path):
    src = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("Hello Heading", level=1)
    doc.add_paragraph("A normal paragraph.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "H1"
    table.cell(0, 1).text = "H2"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "B"
    doc.save(str(src))

    out = tmp_path / "out"
    out.mkdir()
    md = _convert_with_python_docx(src, out)

    assert md is not None
    assert "Hello Heading" in md
    assert "A normal paragraph." in md
    assert "|" in md
    assert "H1" in md
    assert "H2" in md
    assert "| --- | --- |" in md or "| --- |" in md
